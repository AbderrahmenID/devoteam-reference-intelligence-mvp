from __future__ import annotations

import hashlib
import os
import shutil
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Iterable
from zipfile import ZipFile

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from retrieval.schemas import RetrievalResult


ACCENT = "E64B3C"
DARK = "182C3B"
LIGHT = "F3F1EC"
WHITE = "FFFFFF"


@dataclass(frozen=True)
class ExportArtifact:
    path: Path
    reference_count: int
    sha256: str
    structural_checks: dict[str, Any]


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def _shade(cell: Any, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_text(cell: Any, value: str, *, bold: bool = False, color: str = DARK, size: int = 8) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(value)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _add_page_number(paragraph: Any) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Page ")
    run.font.size = Pt(8)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, value, end])


def _configure_document(document: DocumentObject) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    footer = section.footer.paragraphs[0]
    _add_page_number(footer)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9)
    for style_name, size, color in (("Title", 26, ACCENT), ("Heading 1", 20, ACCENT), ("Heading 2", 15, DARK)):
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)


def _display(value: str | None, policy: str) -> str:
    if value:
        return value
    return "Not available in source" if policy == "not_available" else ""


def _add_heading_block(document: DocumentObject, count: int, query: str) -> None:
    title = document.add_heading("Nos principales références", level=1)
    title.paragraph_format.space_after = Pt(4)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    run = subtitle.add_run(f"{count} evidence-gated reference{'s' if count != 1 else ''}")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(DARK)
    if query.strip():
        subtitle.add_run(f" · Search context: {query.strip()}")


def _add_summary_table(document: DocumentObject, results: list[RetrievalResult], policy: str) -> None:
    table = document.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    headers = ("#", "Project / mission", "Client", "Country", "Period", "Key themes")
    widths = (Cm(1.1), Cm(7.8), Cm(4.5), Cm(3.0), Cm(2.4), Cm(7.0))
    for cell, header, width in zip(table.rows[0].cells, headers, widths):
        cell.width = width
        _shade(cell, DARK)
        _set_cell_text(cell, header, bold=True, color=WHITE, size=8)
    for index, result in enumerate(results):
        cells = table.add_row().cells
        values = (
            _display(result.reference_number, policy),
            result.project_title,
            _display(result.client, policy),
            _display(result.country, policy),
            _display(result.period, policy),
            _display("\n".join(f"✓ {theme}" for theme in result.key_themes), policy),
        )
        for cell, value, width in zip(cells, values, widths):
            cell.width = width
            _shade(cell, WHITE if index % 2 == 0 else LIGHT)
            _set_cell_text(cell, value, size=7 if len(value) > 140 else 8)


def _add_detail_table(document: DocumentObject, result: RetrievalResult, policy: str) -> None:
    rows = (
        ("Stable reference ID", result.reference_id),
        ("Mission name", result.mission_name),
        ("Country", result.country),
        ("Contracting authority", result.contracting_authority),
        ("Start date", result.project_start_date or ""),
        ("Completion date", result.completion_date or ""),
        ("Status", result.status or ""),
        ("Sector", result.sector),
        ("Offering", ", ".join(result.offerings)),
        ("Technologies", ", ".join(result.technologies)),
        ("Key themes", ", ".join(result.key_themes)),
        ("Evidence type", ", ".join(result.evidence_types)),
        ("Document languages", ", ".join(result.document_languages)),
    )
    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for label, raw_value in rows:
        cells = table.add_row().cells
        cells[0].width = Cm(5)
        cells[1].width = Cm(12)
        _shade(cells[0], LIGHT)
        _set_cell_text(cells[0], label, bold=True, size=8)
        _set_cell_text(cells[1], _display(raw_value, policy), size=8)


def _add_labelled_text(document: DocumentObject, label: str, value: str, policy: str) -> None:
    heading = document.add_paragraph()
    heading.paragraph_format.space_before = Pt(7)
    heading.paragraph_format.space_after = Pt(2)
    run = heading.add_run(label)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(ACCENT)
    paragraph = document.add_paragraph(_display(value, policy))
    paragraph.paragraph_format.space_after = Pt(4)


def _add_annex(
    document: DocumentObject,
    results: list[RetrievalResult],
    options: dict[str, Any],
    new_section: bool,
) -> None:
    policy = str(options["missing_value_policy"])
    section = document.add_section(WD_SECTION.NEW_PAGE) if new_section else document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    if new_section:
        section.footer.is_linked_to_previous = False
        section.footer.paragraphs[0].clear()
        _add_page_number(section.footer.paragraphs[0])
    for index, result in enumerate(results):
        if index:
            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        number = _display(result.reference_number, policy)
        document.add_heading(f"Référence N°{number}" if number else "Référence", level=2)
        _add_detail_table(document, result, policy)
        _add_labelled_text(document, "Description du projet", result.description, policy)
        if options["include_evidence_passages"]:
            _add_labelled_text(document, "Services rendus et éléments probants", "", "blank")
            for passage in result.supporting_passages:
                quote = document.add_paragraph(style="Intense Quote" if "Intense Quote" in document.styles else None)
                quote.paragraph_format.space_after = Pt(2)
                quote.add_run(passage.text)
                citation = document.add_paragraph()
                citation.paragraph_format.left_indent = Cm(0.5)
                citation.paragraph_format.space_after = Pt(5)
                run = citation.add_run(
                    f"Source: {passage.source_document} · page {passage.source_page} · "
                    f"{passage.citation_label}"
                    + (f" · {passage.citation_uri}" if passage.citation_uri else "")
                )
                run.italic = True
                run.font.size = Pt(7)
                run.font.color.rgb = RGBColor(90, 100, 108)
        if options["include_scores"]:
            scores = result.score_components
            _add_labelled_text(
                document,
                "Retrieval diagnostics",
                (
                    f"Relevance rank {result.relevance_rank}; BM25 {scores.bm25_score:.4f}; "
                    f"dense cosine {scores.dense_cosine:.4f}; hybrid RRF {scores.hybrid_rrf:.6f}."
                ),
                policy,
            )


def validate_docx_structure(path: Path, expected_results: Iterable[RetrievalResult]) -> dict[str, Any]:
    expected = list(expected_results)
    if not path.is_file() or path.stat().st_size == 0:
        raise RuntimeError("DOCX export was not created")
    with ZipFile(path) as archive:
        names = set(archive.namelist())
        required = {"[Content_Types].xml", "word/document.xml", "word/styles.xml"}
        missing = required - names
        if missing:
            raise RuntimeError(f"DOCX package is missing required entries: {sorted(missing)}")
    reopened = Document(path)
    table_text = [
        cell.text
        for table in reopened.tables
        for row in table.rows
        for cell in row.cells
    ]
    full_text = "\n".join(
        [*(paragraph.text for paragraph in reopened.paragraphs), *table_text]
    )
    omitted = [
        result.reference_id
        for result in expected
        if result.reference_id not in full_text and result.project_title not in full_text
    ]
    if omitted:
        raise RuntimeError(f"DOCX export omitted selected references: {omitted}")
    if not reopened.tables:
        raise RuntimeError("DOCX export contains no tables")
    return {
        "package_valid": True,
        "reopened": True,
        "table_count": len(reopened.tables),
        "selected_references_present": len(expected),
    }


def export_docx(
    *,
    template_path: Path,
    output_path: Path,
    results: list[RetrievalResult],
    query: str,
    options: dict[str, Any],
    expected_template_sha256: str,
) -> ExportArtifact:
    if not results:
        raise ValueError("At least one retained reference is required for export")
    source_hash = sha256_file(template_path)
    if source_hash.casefold() != expected_template_sha256.casefold():
        raise RuntimeError("Template hash does not match the configured immutable source")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    temporary_path = output_path.with_suffix(".working.docx")
    template_working_path = output_path.with_suffix(".template-working.docx")
    if temporary_path.exists():
        temporary_path.unlink()
    if template_working_path.exists():
        template_working_path.unlink()
    shutil.copy2(template_path, template_working_path)
    try:
        template_document = Document(template_working_path)
        if len(template_document.tables) < 18:
            raise RuntimeError("Template structure no longer matches the audited summary and annex layout")
        document = Document()
        _configure_document(document)
        if options["include_summary_table"]:
            _add_heading_block(document, len(results), query)
            _add_summary_table(document, results, str(options["missing_value_policy"]))
        if options["include_detailed_annex"]:
            _add_annex(
                document,
                results,
                options,
                new_section=bool(options["include_summary_table"]),
            )
        document.core_properties.title = "Devoteam reference export"
        document.core_properties.subject = "Evidence-gated multilingual reference selection"
        document.core_properties.author = "Devoteam Reference Finder"
        document.core_properties.modified = datetime.now()
        document.save(temporary_path)
        checks = validate_docx_structure(temporary_path, results)
        os.replace(temporary_path, output_path)
    finally:
        if temporary_path.exists():
            temporary_path.unlink()
        if template_working_path.exists():
            template_working_path.unlink()
    if sha256_file(template_path) != source_hash:
        raise RuntimeError("Template source changed during export")
    return ExportArtifact(
        path=output_path,
        reference_count=len(results),
        sha256=sha256_file(output_path),
        structural_checks=checks,
    )
