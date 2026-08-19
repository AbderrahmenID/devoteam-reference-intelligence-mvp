from __future__ import annotations

from pathlib import Path

import pytest
import yaml
from docx import Document

from exporting.docx_export import export_docx, sha256_file, validate_docx_structure
from retrieval.schemas import EvidencePassage, RetrievalResult, ScoreComponents


ROOT = Path(__file__).resolve().parents[1]
CONFIG = yaml.safe_load((ROOT / "config.yaml").read_text(encoding="utf-8"))
TEMPLATE = ROOT / CONFIG["export"]["template_path"]


def _result(reference_id: str, title: str, client: str = "Client source") -> RetrievalResult:
    passage = EvidencePassage(
        text=f"Evidence passage supporting {title}.",
        source_document="source.pdf",
        source_page=3,
        citation_label="source.pdf · p. 3",
        citation_uri="https://example.test/source.pdf",
        language="fr",
    )
    scores = ScoreComponents(
        bm25_score=5.2,
        dense_cosine=0.83,
        hybrid_rrf=0.015,
        query_term_coverage=0.5,
        supporting_passages=1,
    )
    return RetrievalResult(
        reference_id=reference_id,
        reference_number=None,
        project_title=title,
        mission_name=title,
        client=client,
        contracting_authority=client,
        country="Tunisie",
        country_code="TN",
        project_start_date="2021",
        completion_date="2022",
        period="2021–2022",
        status="completed",
        sector="Banque",
        offerings=["PCA/PCI"],
        service_nature=title,
        technologies=[],
        key_themes=["Audit technique et organisationnel"],
        description=title,
        services_delivered=[passage.text],
        supporting_passages=[passage],
        evidence_available=True,
        evidence_types=["ATTESTATION"],
        document_languages=["fr"],
        match_reasons=["Exact terms: pca"],
        rank=1,
        relevance_rank=1,
        score_components=scores,
        title=title,
        offering="PCA/PCI",
        supporting_passage=passage.text,
        source_document=passage.source_document,
        source_page=passage.source_page,
        citation_label=passage.citation_label,
        citation_uri=passage.citation_uri,
        evidence_language="fr",
    )


def _options(**overrides: object) -> dict:
    options = {
        "include_summary_table": True,
        "include_detailed_annex": True,
        "include_evidence_passages": True,
        "include_scores": False,
        "missing_value_policy": "blank",
    }
    options.update(overrides)
    return options


def test_selected_subset_export_reopens_and_preserves_template(tmp_path: Path) -> None:
    source_hash = sha256_file(TEMPLATE)
    selected = [_result("stable-a", "Selected API programme"), _result("stable-b", "Selected cloud programme")]
    output = tmp_path / "selected.docx"
    artifact = export_docx(
        template_path=TEMPLATE,
        output_path=output,
        results=selected,
        query="API cloud",
        options=_options(),
        expected_template_sha256=CONFIG["export"]["template_sha256"],
    )
    assert artifact.reference_count == 2
    assert artifact.structural_checks["table_count"] == 3
    assert artifact.structural_checks["selected_references_present"] == 2
    assert sha256_file(TEMPLATE) == source_hash
    assert validate_docx_structure(output, selected)["reopened"] is True


def test_summary_only_export_contains_only_selected_reference(tmp_path: Path) -> None:
    selected = [_result("stable-selected", "Only selected project", client="")]
    output = tmp_path / "summary-only.docx"
    export_docx(
        template_path=TEMPLATE,
        output_path=output,
        results=selected,
        query="selected",
        options=_options(
            include_detailed_annex=False,
            include_evidence_passages=False,
            missing_value_policy="not_available",
        ),
        expected_template_sha256=CONFIG["export"]["template_sha256"],
    )
    document = Document(output)
    table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    assert len(document.tables) == 1
    assert "Only selected project" in table_text
    assert "Not available in source" in table_text
    assert "Unselected project" not in table_text


def test_export_rejects_empty_selection(tmp_path: Path) -> None:
    with pytest.raises(ValueError, match="At least one"):
        export_docx(
            template_path=TEMPLATE,
            output_path=tmp_path / "empty.docx",
            results=[],
            query="PCA",
            options=_options(),
            expected_template_sha256=CONFIG["export"]["template_sha256"],
        )


def test_export_rejects_a_template_hash_mismatch(tmp_path: Path) -> None:
    with pytest.raises(RuntimeError, match="Template hash"):
        export_docx(
            template_path=TEMPLATE,
            output_path=tmp_path / "bad-hash.docx",
            results=[_result("stable-a", "Project")],
            query="PCA",
            options=_options(),
            expected_template_sha256="0" * 64,
        )


def test_original_and_canonical_template_are_byte_identical() -> None:
    assert sha256_file(ROOT / "templates/Template Ref.docx") == sha256_file(TEMPLATE)


def test_arabic_and_mixed_script_content_survives_export(tmp_path: Path) -> None:
    arabic = _result("stable-ar", "مهمة استمرارية الأعمال")
    arabic.mission_name = "مهمة استمرارية الأعمال للبنوك"
    arabic.description = "دراسة فرنسية وعربية لاستمرارية الأعمال"
    arabic.services_delivered = ["تحليل المخاطر · PCA en Tunisie"]
    arabic.supporting_passages[0].text = "تحليل المخاطر واستمرارية الأعمال · PCA en Tunisie"
    output = tmp_path / "arabic-mixed.docx"
    export_docx(
        template_path=TEMPLATE,
        output_path=output,
        results=[arabic],
        query="PCA للبنوك en Tunisie",
        options=_options(),
        expected_template_sha256=CONFIG["export"]["template_sha256"],
    )
    document = Document(output)
    text = "\n".join(
        [
            *(paragraph.text for paragraph in document.paragraphs),
            *(cell.text for table in document.tables for row in table.rows for cell in row.cells),
        ]
    )
    assert "مهمة استمرارية الأعمال" in text
    assert "PCA en Tunisie" in text
